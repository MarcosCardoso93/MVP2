"""Serviço da HU-14 — arquivo `_ENV` + carta para a operadora.

**Escopo: T-080 (`_ENV`), T-081 (numeração CT), T-082 (renderizar carta), T-083 (salvar
carta).** D-3 foi esclarecida (2026-07-23): o modelo `Base_Contestação_..._M` é o próprio
output da Consolidação (`montar_contest`/T-023 + `consolidar_expectativa_vivo`/T-021).

**T-082/T-083 (2026-07-27):** o modelo de carta foi substancialmente resolvido a partir
de dois exemplos reais já emitidos (`DOCS/Outros arquivos auxiliares/CT 251-2026...` e
`CT 252-2026...`, cenário SEM RETENÇÃO — ver D-3/D-13 em `TODO/decisoes.md`). Achado
importante: a tabela inserida na carta **é** quebrada por remuneração (corrige a nota
antiga de D-13, que assumia o contrário sem ver um exemplo real). **Assinatura (D-3,
resolvida em 2026-07-27):** os dois exemplos reais tinham signatários diferentes; o
usuário confirmou usar sempre a assinatura da carta real CT 252-2026 (Angélica) —
`ProvedorAssinaturaCartaPadrao`, sem mais placeholder.
"""

from __future__ import annotations

import re
import shutil
from datetime import date
from pathlib import Path
from typing import Callable, Optional

import pandas as pd
from docx import Document

from src.config import constantes_epico4 as const
from src.config.logger_config import logger
from src.services import mapa_remuneracao as mr
from src.utils import estrutura_pastas as ep
from src.utils import gerenciador_arquivos as ga
from src.utils import nomenclatura as nom
from src.utils.decoradores import log_execucao

# Nomes das abas mantidas no _ENV (AI/09 §4.1: "manter apenas as abas Contest e TBRA").
ABA_CONTEST = const.ABA_CONTEST
ABA_TBRA = const.ABA_TBRA

# Colunas da tabela por remuneração inserida na carta (T-082) — sem
# `contestacao_a_enviar`/"Contestar" (AI/09 §4.2 item 4).
_COLUNAS_TABELA_CARTA = [
    "EOT",
    "minutos_tbra",
    "vb_tbra",
    "minutos_operadora",
    "vb_operadora",
    "minutos_diferenca",
    "vb_diferenca",
    "minutos_variacao_perc",
    "vb_variacao_perc",
]

_MES_EXTENSO = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
}

# Casa "CT 362", "CT-362", "CT - 362", "CT_362..." (com ou sem espaços/traço),
# case-insensitive — tolerante a variações de grafia observadas no nome real
# ("CT- 334/2025/WA*ID03", ver AI/09 §4.2 e o screenshot da carta CT-334/2025).
_PADRAO_NUMERO_CT = re.compile(r"CT[\s_-]*(\d+)", re.IGNORECASE)


def _extrair_numero_ct(nome_arquivo: str) -> int | None:
    """Extrai o número da carta de um nome de arquivo/pasta, ou `None` se não casar."""

    correspondencia = _PADRAO_NUMERO_CT.search(nome_arquivo)
    if not correspondencia:
        return None
    return int(correspondencia.group(1))


@log_execucao
def obter_proximo_numero_carta(pasta_controle_ct: Path) -> int:
    """
    Determina o próximo número sequencial de carta CT (T-081).

    Lê os nomes de arquivo/pasta em `pasta_controle_ct` (tipicamente
    `estrutura_pastas.caminho_controle_ct(aaaamm)`), extrai o maior número `CT` encontrado
    e retorna o próximo (AI/09 §4.2: "último `CT 362...` ⇒ novo `CT - 363...`").

    Args:
        pasta_controle_ct: Pasta `...\\Correspondências Enviadas\\CT\\{ano}`.

    Returns:
        O próximo número sequencial. Se a pasta não existir ou estiver vazia (ou sem
        nenhum nome reconhecível como `CT <n>`), retorna `1` (bootstrap — não é regra de
        negócio documentada, apenas o ponto de partida seguro de uma sequência nova).
    """

    pasta_controle_ct = Path(pasta_controle_ct)

    if not pasta_controle_ct.is_dir():
        logger.warning(
            f"[HU-14] Pasta de controle CT inexistente [{pasta_controle_ct}] — "
            f"assumindo primeira carta (nº 1)."
        )
        return 1

    numeros = [
        numero
        for caminho in pasta_controle_ct.iterdir()
        if (numero := _extrair_numero_ct(caminho.name)) is not None
    ]

    if not numeros:
        logger.info(f"[HU-14] Nenhuma carta CT encontrada em [{pasta_controle_ct}] — nº 1.")
        return 1

    proximo = max(numeros) + 1
    logger.info(f"[HU-14] Último número CT encontrado: {max(numeros)}. Próximo: {proximo}.")
    return proximo


def nome_proxima_carta(pasta_controle_ct: Path) -> str:
    """Atalho: resolve o próximo número (T-081) e monta o nome `CT - {n}` (T-006)."""

    numero = obter_proximo_numero_carta(pasta_controle_ct)
    return nom.nome_carta(numero)


@log_execucao
def montar_abas_env(
    df_contest: pd.DataFrame,
    df_tbra_bruta: pd.DataFrame,
    referencia: str,
    indice_descritor: int,
    indice_remuneracao: dict[str, list[str]],
    obter_tipo_contestacao: Callable[[str, str, str, str, str], Optional[str]],
    indice_credora: int = const.COL_CREDORA,
    indice_devedora: int = const.COL_DEVEDORA,
    indice_trafego: int = const.COL_TRAFEGO,
) -> dict[str, pd.DataFrame]:
    """
    Monta as abas `Contest` e `TBRA` do `_ENV` (T-080), só com as linhas contestadas.

    AI/09 §4.1: "manter apenas as abas Contest e TBRA; apagar as demais" e "remover as
    linhas do que não será contestado — deixar somente os dados contestados". Como o
    modelo `_M` foi esclarecido como sendo o próprio output da Consolidação (D-3), as
    abas `RESUMO`/`{operadora}`/`DE PARA EOT` já são descartadas por não fazerem parte do
    conjunto de entrada desta função — não é necessário "apagá-las" explicitamente.

    **Gate de "linha contestada"** — mesmo critério de `geracao_cont_proc` (D-13/T-100):
    `contestacao_a_enviar == "S"` **e** sinal do analista não `None` (cenário SEM/COM
    retenção; nunca "sem contestação").

    Args:
        df_contest: Saída de `consolidacao_contestacao.montar_contest` (T-023).
        df_tbra_bruta: Aba `TBRA` consolidada (T-021), linha a linha (sem agregação),
            até `R$_Bruto`.
        referencia: Mês de referência do Detraf da contestação (`AAAAMM`).
        indice_descritor: Índice (0-based) da coluna de descritor em `df_tbra_bruta` —
            sem posição fixa na documentação, informado pelo chamador.
        indice_remuneracao: Índice `descritor final -> [remunerações]` (D-5).
        obter_tipo_contestacao: Callable injetada (repository via controller) que resolve
            o sinal COM/SEM retenção por
            `(eot_operadora, eot_tbra, referencia, trafego, remuneracao)` (T-024/D-6; a
            remuneração entra na chave desde 2026-07-28).
        indice_credora, indice_devedora, indice_trafego: Índices documentados (D-8) usados
            para casar cada linha bruta da `TBRA` com a chave contestada do `Contest`.

    Returns:
        Dicionário ``{"Contest": DataFrame, "TBRA": DataFrame}``. Ambos vazios se nada
        for contestado.
    """

    if df_contest.empty:
        return {ABA_CONTEST: df_contest.copy(), ABA_TBRA: df_tbra_bruta.iloc[0:0].copy()}

    candidatas = df_contest[df_contest["contestacao_a_enviar"] == "S"].copy()

    if not candidatas.empty:
        candidatas["_sinal"] = candidatas.apply(
            lambda linha: obter_tipo_contestacao(
                linha["eot_credora"], linha["eot_devedora"], referencia, linha["trafego"],
                linha["remuneracao"],
            ),
            axis=1,
        )
        contestadas = candidatas[candidatas["_sinal"].notna()].drop(columns=["_sinal"])
    else:
        contestadas = candidatas

    if contestadas.empty:
        logger.info("[HU-14 _ENV] Nenhuma linha contestada — abas ficam vazias.")
        return {
            ABA_CONTEST: df_contest.iloc[0:0].copy(),
            ABA_TBRA: df_tbra_bruta.iloc[0:0].copy(),
        }

    chaves_contestadas = set(
        zip(
            contestadas["eot_credora"],
            contestadas["eot_devedora"],
            contestadas["remuneracao"],
            contestadas["trafego"],
        )
    )

    tbra = df_tbra_bruta.copy()
    tbra["_remuneracao"] = mr.resolver_series(
        tbra.iloc[:, indice_descritor], indice_remuneracao
    )
    tbra["_chave"] = list(
        zip(
            tbra.iloc[:, indice_credora].astype(str).str.strip(),
            tbra.iloc[:, indice_devedora].astype(str).str.strip(),
            tbra["_remuneracao"],
            tbra.iloc[:, indice_trafego].astype(str).str.strip(),
        )
    )
    tbra_filtrada = tbra[tbra["_chave"].isin(chaves_contestadas)].drop(
        columns=["_remuneracao", "_chave"]
    )

    logger.info(
        f"[HU-14 _ENV] Contest: {contestadas.shape[0]} linha(s) contestada(s); "
        f"TBRA: {tbra_filtrada.shape[0]} linha(s) correspondente(s)."
    )
    return {ABA_CONTEST: contestadas, ABA_TBRA: tbra_filtrada}


@log_execucao
def gerar_arquivo_env(
    abas_env: dict[str, pd.DataFrame],
    operadora: str,
    aaaamm: str,
    raiz_operadoras: Path | None = None,
) -> Optional[Path]:
    """
    Grava o arquivo `_ENV` na pasta `Contestações` (T-080) — **só se houver contestação**.

    Se ambas as abas estiverem vazias (nenhuma linha contestada), **nenhum arquivo é
    criado** — mesmo critério de `gerar_arquivo_int`/`gerar_arquivo_cont_proc`: o `_ENV`
    só existe nos cenários SEM/COM retenção (AI/09 §0, matriz).

    Args:
        abas_env: Saída de :func:`montar_abas_env` (``{"Contest": df, "TBRA": df}``).
        operadora: Nome da operadora (usado no nome do arquivo e na pasta).
        aaaamm: Período de referência (`AAAAMM`).
        raiz_operadoras: Raiz `...\\Operadoras`. Default: `CAMINHO_OPERADORAS` (config).

    Returns:
        Caminho do arquivo `.xlsx` gravado, ou ``None`` se nada foi contestado.
    """

    if all(df.empty for df in abas_env.values()):
        logger.info("[HU-14 _ENV] Nada a gravar (nenhuma linha contestada).")
        return None

    pasta_contestacoes = ep.caminho_contestacoes(
        operadora, aaaamm, raiz_operadoras=raiz_operadoras, criar=True
    )
    nome_arquivo = nom.nome_env(operadora, aaaamm)
    caminho = pasta_contestacoes / f"{nome_arquivo}{const.EXTENSAO_EXCEL}"

    ga.salvar_planilhas(abas_env, caminho, incluir_cabecalho=True)

    logger.info(f"[HU-14 _ENV] Arquivo gravado em [{caminho}].")
    return caminho


class ProvedorAssinaturaCarta:
    """Contrato para resolver o signatário (nome, cargo) da carta.

    Isolado atrás desta interface porque os dois exemplos reais (CT 251/CT 252-2026,
    mesmo dia) tinham signatários diferentes com o mesmo cargo genérico — a escolha não
    era óbvia até o usuário decidir (D-3, resolvida em 2026-07-27).
    """

    def resolver(self, operadora: str) -> tuple[str, str]:
        raise NotImplementedError


class ProvedorAssinaturaCartaPadrao(ProvedorAssinaturaCarta):
    """Assinatura fixa confirmada pelo usuário (D-3, 2026-07-27).

    Usa sempre a assinatura da carta real
    ``CT 252-2026-DIR-A1-tbrasilxserctel-Contestação_202606.docx`` (Angélica), para
    qualquer operadora — não há distinção por operadora/região.
    """

    def resolver(self, operadora: str) -> tuple[str, str]:
        return (const.CARTA_ASSINATURA_NOME, const.CARTA_ASSINATURA_CARGO)


@log_execucao
def montar_tabelas_carta(df_contest_contestado: pd.DataFrame) -> dict[str, dict]:
    """
    Monta as tabelas por remuneração inseridas no corpo da carta (T-082).

    **Granularidade (corrige D-13, 2026-07-27):** os exemplos reais (CT 251/252-2026)
    confirmam que a tabela da carta é quebrada **por remuneração** (blocos "TU-RL",
    "VU-M" etc., cada um com sub-rótulo de tipo de serviço quando disponível) — não há
    colapso por EOT×Tráfego como se assumia antes. Usa a aba `Contest` já filtrada por
    `montar_abas_env` (T-080) diretamente, sem agregação adicional.

    Args:
        df_contest_contestado: Aba `Contest` já filtrada (saída de
            ``montar_abas_env(...)["Contest"]``, T-080) — colunas `eot_credora`,
            `eot_devedora`, `remuneracao`, `minutos_tbra`, `vb_tbra`,
            `minutos_operadora`, `vb_operadora`, `minutos_diferenca`, `vb_diferenca`,
            `minutos_variacao_perc`, `vb_variacao_perc` e, opcionalmente,
            `tipo_operacao` (rótulo SMS/STFC visto nos exemplos reais).

    Returns:
        Dicionário ``{remuneracao: {"tipo_operacao": str | None, "tabela": DataFrame}}``,
        na ordem em que a remuneração aparece em ``df_contest_contestado``. Cada
        `DataFrame` traz a coluna `EOT` (formato `devedora/credora`, D-11) + as colunas
        numéricas, com uma linha `TOTAL` ao final (variação percentual em branco na
        linha `TOTAL` — mesmo padrão visto nos exemplos reais). **Sem** a coluna
        `contestacao_a_enviar`/"Contestar" (AI/09 §4.2). Vazio se não houver nada
        contestado.
    """

    if df_contest_contestado.empty:
        return {}

    df = df_contest_contestado.copy()
    df["EOT"] = (
        df["eot_devedora"].astype(str) + "/" + df["eot_credora"].astype(str)
    )

    resultado: dict[str, dict] = {}
    for remuneracao, grupo in df.groupby("remuneracao", sort=False):
        tipo_operacao = None
        if "tipo_operacao" in grupo.columns:
            valores = grupo["tipo_operacao"].dropna().unique()
            if len(valores) == 1:
                tipo_operacao = valores[0]

        tabela = grupo[_COLUNAS_TABELA_CARTA].reset_index(drop=True)

        totais = tabela.drop(columns="EOT").sum(numeric_only=True).to_dict()
        linha_total = {"EOT": "TOTAL", **totais}
        # Variação percentual não é somável — linha TOTAL fica em branco (mesmo padrão
        # visto nos exemplos reais CT 251/252-2026).
        linha_total["minutos_variacao_perc"] = None
        linha_total["vb_variacao_perc"] = None

        tabela = pd.concat([tabela, pd.DataFrame([linha_total])], ignore_index=True)
        resultado[str(remuneracao)] = {"tipo_operacao": tipo_operacao, "tabela": tabela}

    logger.info(f"[HU-14 Carta] {len(resultado)} tabela(s) de remuneração montada(s).")
    return resultado


def _inserir_tabela(documento: Document, tabela: pd.DataFrame) -> None:
    """Insere `tabela` como uma tabela nativa do Word em `documento` (T-082)."""

    colunas = list(tabela.columns)
    tabela_docx = documento.add_table(rows=1, cols=len(colunas))
    for indice, coluna in enumerate(colunas):
        tabela_docx.rows[0].cells[indice].text = coluna

    for _, linha in tabela.iterrows():
        celulas = tabela_docx.add_row().cells
        for indice, coluna in enumerate(colunas):
            valor = linha[coluna]
            celulas[indice].text = "" if pd.isna(valor) else str(valor)


@log_execucao
def renderizar_carta(
    numero_ct: int,
    data_carta: date,
    aaaamm: str,
    operadora: str,
    tipo_contestacao: str,
    tabelas_por_remuneracao: dict[str, dict],
    cidade: Optional[str] = None,
    provedor_assinatura: Optional[ProvedorAssinaturaCarta] = None,
) -> Document:
    """
    Renderiza o `.docx` da carta de contestação (T-082), a partir dos exemplos reais
    CT 251/252-2026 (D-3).

    **Não replica** logo/rodapé/fontes do arquivo real — mesmo gap aceito em D-12
    (EXT/INT): conteúdo é o critério de aceite testável, formatação exata fica para
    quando houver necessidade real.

    Args:
        numero_ct: Número sequencial da carta (T-081, `obter_proximo_numero_carta`).
        data_carta: Data de emissão (injetada pelo chamador — nunca `datetime.now()`
            dentro do service, para manter a função determinística/testável).
        aaaamm: Mês de referência do Detraf (usado no "Assunto:" e no corpo).
        operadora: Nome da operadora (usado no corpo).
        tipo_contestacao: `constantes_epico4.CENARIO_SEM_RETENCAO` ou
            `CENARIO_COM_RETENCAO` (texto COM retenção é assumido por substituição
            simples da frase SEM retenção — sem exemplo real ainda, ver D-3).
        tabelas_por_remuneracao: Saída de :func:`montar_tabelas_carta`.
        cidade: Cidade da data de emissão. **RESOLVIDO (D-3, 2026-07-27):** usuário
            confirmou `constantes_epico4.CARTA_CIDADE_PADRAO` ("São Paulo") como padrão
            fixo — os 2 exemplos reais divergiam no mesmo dia (Rio de Janeiro/São Paulo).
            Parâmetro mantido para permitir override pontual, se necessário.
        provedor_assinatura: Resolve `(nome, cargo)` do signatário; default
            :class:`ProvedorAssinaturaCartaPadrao` (assinatura fixa confirmada em D-3).

    Returns:
        Documento `python-docx` pronto para ser salvo (:func:`gerar_arquivo_carta`).
    """

    if cidade is None:
        cidade = const.CARTA_CIDADE_PADRAO

    documento = Document()

    documento.add_paragraph(f"{const.PREFIXO_CARTA}- {numero_ct}/{data_carta.year}")
    documento.add_paragraph(
        f"{cidade}, {data_carta.day} de {_MES_EXTENSO[data_carta.month]} de "
        f"{data_carta.year}."
    )
    documento.add_paragraph(
        f"ASSUNTO: {const.CARTA_ASSUNTO_TEMPLATE.format(aaaamm=aaaamm)}"
    )
    documento.add_paragraph(const.CARTA_SAUDACAO)
    documento.add_paragraph(
        const.CARTA_CORPO_TEMPLATE.format(
            tipo_contestacao=tipo_contestacao,
            aaaamm=aaaamm,
            operadora=operadora.upper(),
        )
    )

    for remuneracao, dados in tabelas_por_remuneracao.items():
        titulo = remuneracao
        if dados.get("tipo_operacao"):
            titulo = f"{remuneracao}    {dados['tipo_operacao']}"
        documento.add_paragraph(titulo)
        _inserir_tabela(documento, dados["tabela"])

    documento.add_paragraph(const.CARTA_FECHO)

    provedor = provedor_assinatura or ProvedorAssinaturaCartaPadrao()
    nome, cargo = provedor.resolver(operadora)
    documento.add_paragraph(nome)
    documento.add_paragraph(cargo)

    logger.info(
        f"[HU-14 Carta] Carta CT-{numero_ct} renderizada "
        f"({len(tabelas_por_remuneracao)} tabela(s))."
    )
    return documento


@log_execucao
def gerar_arquivo_carta(
    documento: Document,
    operadora: str,
    numero_ct: int,
    aaaamm: str,
    raiz_operadoras: Path | None = None,
    raiz_controle_ct: Path | None = None,
) -> Path:
    """
    Salva a carta (T-083) em `Contestações` e uma cópia em `...\\CT\\{ano}` (AI/09 §4.2
    passo 5).

    Args:
        documento: Saída de :func:`renderizar_carta`.
        operadora: Nome da operadora (usado na pasta).
        numero_ct: Número sequencial da carta (usado no nome do arquivo).
        aaaamm: Período de referência (usado na pasta/ano de controle).
        raiz_operadoras: Raiz `...\\Operadoras`. Default: `CAMINHO_OPERADORAS` (config).
        raiz_controle_ct: Raiz `...\\Correspondências Enviadas\\CT`. Default:
            `CAMINHO_CONTROLE_CT` (config).

    Returns:
        Caminho do arquivo `.docx` gravado em `Contestações`.
    """

    pasta_contestacoes = ep.caminho_contestacoes(
        operadora, aaaamm, raiz_operadoras=raiz_operadoras, criar=True
    )
    nome_arquivo = f"{nom.nome_carta(numero_ct)}{const.EXTENSAO_DOCX}"
    caminho = pasta_contestacoes / nome_arquivo
    documento.save(str(caminho))

    pasta_controle = ep.caminho_controle_ct(
        aaaamm, raiz_controle_ct=raiz_controle_ct, criar=True
    )
    caminho_copia = pasta_controle / nome_arquivo
    shutil.copy2(caminho, caminho_copia)

    logger.info(
        f"[HU-14 Carta] Carta gravada em [{caminho}], cópia em [{caminho_copia}]."
    )
    return caminho
