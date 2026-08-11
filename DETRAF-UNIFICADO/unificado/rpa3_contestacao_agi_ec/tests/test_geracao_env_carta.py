"""Testes da HU-14: numeração CT (T-081), escrita do `_ENV` (T-080) e da carta
(T-082/T-083, a partir dos exemplos reais CT 251/252-2026 — D-3). Assinatura fixa
(Angélica, CT 252-2026) via `ProvedorAssinaturaCartaPadrao`.
"""

from datetime import date
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

from comum.config import constantes as const
from src.services import geracao_env_carta as gec
from src.services import mapa_remuneracao as mr


# ---------------------------------------------------------------------------
# Numeração CT
#
# MUDANÇA DE COMPORTAMENTO (decisão do cliente, 2026-07-31): não pode haver
# duplicação de número. Quando o último número não puder ser identificado, o
# robô ACUSA O ERRO em vez de devolver 1 — a versão anterior reemitiria a carta
# nº 1 sobre uma sequência já existente sempre que a pasta estivesse
# inacessível.
# ---------------------------------------------------------------------------


def test_pasta_inexistente_levanta_erro(tmp_path: Path):
    with pytest.raises(gec.NumeracaoCartaIndeterminada, match="não encontrada"):
        gec.obter_proximo_numero_carta(tmp_path / "nao_existe")


def test_pasta_vazia_levanta_erro(tmp_path: Path):
    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)

    with pytest.raises(gec.NumeracaoCartaIndeterminada, match="CT_NUMERO_INICIAL"):
        gec.obter_proximo_numero_carta(pasta)


def test_pasta_sem_carta_reconhecivel_levanta_erro(tmp_path: Path):
    """Pasta com conteúdo, mas nenhum nome que pareça uma carta CT."""
    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)
    (pasta / "leia-me.txt").write_text("x")
    (pasta / "planilha_de_controle.xlsx").write_text("x")

    with pytest.raises(gec.NumeracaoCartaIndeterminada):
        gec.obter_proximo_numero_carta(pasta)


def test_pasta_vazia_com_numero_inicial_configurado(tmp_path: Path, monkeypatch):
    """
    A exceção deliberada: em janeiro a pasta do ano é legitimamente nova. Com
    `CT_NUMERO_INICIAL` definido, a sequência começa dali — nunca por default.
    """
    from comum.config import configuration

    monkeypatch.setattr(configuration, "CT_NUMERO_INICIAL", 500)

    pasta = tmp_path / "CT" / "2026"
    pasta.mkdir(parents=True)

    assert gec.obter_proximo_numero_carta(pasta) == 500


def test_numero_inicial_nao_sobrepoe_sequencia_existente(tmp_path: Path, monkeypatch):
    """`CT_NUMERO_INICIAL` só vale quando não há carta nenhuma na pasta."""
    from comum.config import configuration

    monkeypatch.setattr(configuration, "CT_NUMERO_INICIAL", 500)

    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)
    (pasta / "CT - 362_CLARO_202507.docx").write_text("x")

    assert gec.obter_proximo_numero_carta(pasta) == 363


def test_encontra_maior_numero_e_soma_um(tmp_path: Path):
    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)
    (pasta / "CT - 340_AMPERNET_202506.docx").write_text("x")
    (pasta / "CT - 362_CLARO_202507.docx").write_text("x")

    assert gec.obter_proximo_numero_carta(pasta) == 363


def test_tolera_variacoes_de_grafia(tmp_path: Path):
    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)
    (pasta / "CT334_AMPERNET.docx").write_text("x")
    (pasta / "CT-334_AMPERNET.docx").write_text("x")
    (pasta / "CT_334_AMPERNET.docx").write_text("x")

    # Todas casam com o mesmo número (334) — próximo = 335.
    assert gec.obter_proximo_numero_carta(pasta) == 335


def test_ignora_arquivos_sem_padrao_ct(tmp_path: Path):
    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)
    (pasta / "CT - 100_CLARO.docx").write_text("x")
    (pasta / "outro_arquivo_qualquer.txt").write_text("x")

    assert gec.obter_proximo_numero_carta(pasta) == 101


def test_nome_proxima_carta(tmp_path: Path):
    pasta = tmp_path / "CT" / "2025"
    pasta.mkdir(parents=True)
    (pasta / "CT - 362_CLARO.docx").write_text("x")

    assert gec.nome_proxima_carta(pasta) == "CT - 363"


# --------------------------------------------------------------------------
# T-080 — escrita do `_ENV`
# --------------------------------------------------------------------------
INDICE_DESCRITOR_TESTE = 6


class _SinalFake:
    def __init__(self, respostas: dict[tuple, str | None]):
        self._respostas = respostas

    def __call__(self, eot_operadora, eot_tbra, referencia, trafego, remuneracao):
        return self._respostas.get((eot_operadora, eot_tbra, referencia, trafego, remuneracao))


def _linha_tbra(credora, devedora, trafego, descritor, minutos="10,0", r_bruto="1,00"):
    linha = ["0"] * 15
    linha[0] = credora
    linha[1] = devedora
    linha[2] = "202507"
    linha[3] = trafego
    linha[6] = descritor
    linha[9] = minutos
    linha[14] = r_bruto
    return linha


def _contest_df(linhas: list[dict]) -> pd.DataFrame:
    return pd.DataFrame(linhas)


def _linha_contest(
    eot_credora,
    eot_devedora,
    trafego,
    remuneracao="TU-RL",
    contestacao_a_enviar="S",
):
    return {
        "eot_credora": eot_credora,
        "eot_devedora": eot_devedora,
        "remuneracao": remuneracao,
        "trafego": trafego,
        "contestacao_a_enviar": contestacao_a_enviar,
    }


def test_montar_abas_env_vazio_quando_sem_contestacao(indice_remuneracao):
    df_contest = _contest_df([_linha_contest("200", "011", "202507", contestacao_a_enviar="N")])
    df_tbra = pd.DataFrame([_linha_tbra("200", "011", "202507", "XPTO_L")])

    abas = gec.montar_abas_env(
        df_contest, df_tbra, "202507", INDICE_DESCRITOR_TESTE, indice_remuneracao, _SinalFake({})
    )

    assert abas[gec.ABA_CONTEST].empty
    assert abas[gec.ABA_TBRA].empty


def test_montar_abas_env_exclui_sem_sinal_do_analista(indice_remuneracao):
    # Flag S mas analista ainda não decidiu (sinal None) => não entra (mesmo gate do CONT_PROC).
    df_contest = _contest_df([_linha_contest("200", "011", "202507")])
    df_tbra = pd.DataFrame([_linha_tbra("200", "011", "202507", "XPTO_L")])

    abas = gec.montar_abas_env(
        df_contest, df_tbra, "202507", INDICE_DESCRITOR_TESTE, indice_remuneracao, _SinalFake({})
    )

    assert abas[gec.ABA_CONTEST].empty
    assert abas[gec.ABA_TBRA].empty


def test_montar_abas_env_filtra_contest_e_tbra_corretamente(indice_remuneracao):
    df_contest = _contest_df(
        [
            _linha_contest("200", "011", "202507", remuneracao="TU-RL"),
            _linha_contest("201", "012", "202507", remuneracao="VU-M", contestacao_a_enviar="N"),
        ]
    )
    df_tbra = pd.DataFrame(
        [
            _linha_tbra("200", "011", "202507", "XPTO_L"),  # remuneracao TU-RL -> contestada
            _linha_tbra("201", "012", "202507", "XPTO_V"),  # remuneracao VU-M -> nao contestada
        ]
    )
    sinal = _SinalFake({("200", "011", "202507", "202507", "TU-RL"): "com retenção"})

    abas = gec.montar_abas_env(
        df_contest, df_tbra, "202507", INDICE_DESCRITOR_TESTE, indice_remuneracao, sinal
    )

    assert abas[gec.ABA_CONTEST].shape[0] == 1
    assert abas[gec.ABA_CONTEST].iloc[0]["eot_credora"] == "200"
    assert abas[gec.ABA_TBRA].shape[0] == 1
    assert abas[gec.ABA_TBRA].iloc[0, 0] == "200"


def test_gerar_arquivo_env_nao_grava_quando_vazio(tmp_path: Path):
    abas_vazias = {gec.ABA_CONTEST: pd.DataFrame(), gec.ABA_TBRA: pd.DataFrame()}
    caminho = gec.gerar_arquivo_env(
        abas_vazias, operadora="Claro", aaaamm="202507", raiz_operadoras=tmp_path
    )
    assert caminho is None
    assert not any(tmp_path.rglob("*.xlsx"))


def test_gerar_arquivo_env_grava_no_caminho_e_nome_corretos(tmp_path: Path, indice_remuneracao):
    df_contest = _contest_df([_linha_contest("200", "011", "202507")])
    df_tbra = pd.DataFrame([_linha_tbra("200", "011", "202507", "XPTO_L")])
    sinal = _SinalFake({("200", "011", "202507", "202507", "TU-RL"): "com retenção"})

    abas = gec.montar_abas_env(
        df_contest, df_tbra, "202507", INDICE_DESCRITOR_TESTE, indice_remuneracao, sinal
    )
    caminho = gec.gerar_arquivo_env(
        abas, operadora="Claro", aaaamm="202507", raiz_operadoras=tmp_path
    )

    assert caminho.name == "Base Contestação_Claro_202507_ENV.xlsx"
    assert caminho.parent == tmp_path / "Claro" / "2025" / "202507" / "Contestações"
    assert caminho.is_file()

    # Confere que as duas abas foram gravadas.
    abas_recarregadas = pd.read_excel(caminho, engine="openpyxl", sheet_name=None)
    assert set(abas_recarregadas.keys()) == {gec.ABA_CONTEST, gec.ABA_TBRA}


# --------------------------------------------------------------------------
# T-082 — montar_tabelas_carta (agrupamento por remuneração + linha TOTAL)
# --------------------------------------------------------------------------
def _linha_contest_completa(
    eot_credora,
    eot_devedora,
    remuneracao="TU-RL",
    minutos_tbra=0.0,
    vb_tbra=0.0,
    minutos_operadora=0.0,
    vb_operadora=0.0,
    minutos_diferenca=0.0,
    vb_diferenca=0.0,
    minutos_variacao_perc=100.0,
    vb_variacao_perc=100.0,
    tipo_operacao=None,
):
    linha = {
        "eot_credora": eot_credora,
        "eot_devedora": eot_devedora,
        "remuneracao": remuneracao,
        "minutos_tbra": minutos_tbra,
        "vb_tbra": vb_tbra,
        "minutos_operadora": minutos_operadora,
        "vb_operadora": vb_operadora,
        "minutos_diferenca": minutos_diferenca,
        "vb_diferenca": vb_diferenca,
        "minutos_variacao_perc": minutos_variacao_perc,
        "vb_variacao_perc": vb_variacao_perc,
    }
    if tipo_operacao is not None:
        linha["tipo_operacao"] = tipo_operacao
    return linha


def test_montar_tabelas_carta_vazio():
    assert gec.montar_tabelas_carta(pd.DataFrame()) == {}


def test_montar_tabelas_carta_agrupa_por_remuneracao():
    df = pd.DataFrame(
        [
            _linha_contest_completa("200", "011", remuneracao="TU-RL", vb_operadora=0.1),
            _linha_contest_completa("200", "011", remuneracao="VU-M", vb_operadora=125.58),
        ]
    )

    resultado = gec.montar_tabelas_carta(df)

    assert set(resultado.keys()) == {"TU-RL", "VU-M"}
    assert resultado["TU-RL"]["tabela"].shape[0] == 2  # 1 linha + TOTAL
    assert resultado["TU-RL"]["tabela"].iloc[0]["EOT"] == "011/200"


def test_montar_tabelas_carta_linha_total_soma_valores_numericos():
    df = pd.DataFrame(
        [
            _linha_contest_completa("200", "011", vb_operadora=0.1, minutos_diferenca=10),
            _linha_contest_completa("201", "011", vb_operadora=31.7, minutos_diferenca=20),
        ]
    )

    resultado = gec.montar_tabelas_carta(df)
    total = resultado["TU-RL"]["tabela"].iloc[-1]

    assert total["EOT"] == "TOTAL"
    assert total["vb_operadora"] == pytest.approx(31.8)
    assert total["minutos_diferenca"] == pytest.approx(30)


def test_montar_tabelas_carta_variacao_perc_em_branco_no_total():
    df = pd.DataFrame([_linha_contest_completa("200", "011")])

    resultado = gec.montar_tabelas_carta(df)
    total = resultado["TU-RL"]["tabela"].iloc[-1]

    assert pd.isna(total["minutos_variacao_perc"])
    assert pd.isna(total["vb_variacao_perc"])


def test_montar_tabelas_carta_nao_inclui_coluna_contestacao_a_enviar():
    df = pd.DataFrame([{**_linha_contest_completa("200", "011"), "contestacao_a_enviar": "S"}])

    resultado = gec.montar_tabelas_carta(df)

    assert "contestacao_a_enviar" not in resultado["TU-RL"]["tabela"].columns


def test_montar_tabelas_carta_expoe_tipo_operacao_quando_uniforme():
    df = pd.DataFrame(
        [_linha_contest_completa("200", "011", tipo_operacao=const.TIPO_OPERACAO_SMP)]
    )

    resultado = gec.montar_tabelas_carta(df)

    assert resultado["TU-RL"]["tipo_operacao"] == const.TIPO_OPERACAO_SMP


# --------------------------------------------------------------------------
# T-082 — renderizar_carta / assinatura (D-3 resolvida: fixa, Angélica/CT 252-2026)
# --------------------------------------------------------------------------
def test_provedor_assinatura_base_e_abstrato():
    with pytest.raises(NotImplementedError):
        gec.ProvedorAssinaturaCarta().resolver("Claro")


def test_provedor_assinatura_padrao_usa_assinatura_fixa_confirmada():
    nome, cargo = gec.ProvedorAssinaturaCartaPadrao().resolver("Claro")
    assert nome == const.CARTA_ASSINATURA_NOME
    assert cargo == const.CARTA_ASSINATURA_CARGO


def test_provedor_assinatura_padrao_e_igual_para_operadoras_diferentes():
    # Assinatura fixa (D-3, 2026-07-27): não varia por operadora.
    provedor = gec.ProvedorAssinaturaCartaPadrao()
    assert provedor.resolver("Claro") == provedor.resolver("Sercomtel")


def test_renderizar_carta_conteudo_basico():
    tabelas = gec.montar_tabelas_carta(
        pd.DataFrame([_linha_contest_completa("200", "011", remuneracao="TU-RL")])
    )

    documento = gec.renderizar_carta(
        numero_ct=363,
        data_carta=date(2026, 7, 20),
        aaaamm="202606",
        operadora="Claro",
        tipo_contestacao=const.CENARIO_SEM_RETENCAO,
        tabelas_por_remuneracao=tabelas,
        cidade="Rio de Janeiro",
    )

    texto = "\n".join(p.text for p in documento.paragraphs)
    assert "CT- 363/2026" in texto
    assert "Rio de Janeiro, 20 de Julho de 2026." in texto
    assert "ASSUNTO: CONTESTAÇÃO DETRAF – 202606" in texto
    assert "SEM retenção" in texto
    assert len(documento.tables) == 1


def test_renderizar_carta_texto_com_retencao():
    tabelas = gec.montar_tabelas_carta(
        pd.DataFrame([_linha_contest_completa("200", "011")])
    )

    documento = gec.renderizar_carta(
        numero_ct=1,
        data_carta=date(2026, 7, 20),
        aaaamm="202606",
        operadora="Claro",
        tipo_contestacao=const.CENARIO_COM_RETENCAO,
        tabelas_por_remuneracao=tabelas,
        cidade="Rio de Janeiro",
    )

    texto = "\n".join(p.text for p in documento.paragraphs)
    assert "COM retenção" in texto


def test_renderizar_carta_sem_cidade_usa_padrao_sao_paulo():
    # D-3 (2026-07-27): usuário confirmou São Paulo como cidade padrão da data.
    tabelas = gec.montar_tabelas_carta(
        pd.DataFrame([_linha_contest_completa("200", "011")])
    )

    documento = gec.renderizar_carta(
        numero_ct=1,
        data_carta=date(2026, 7, 20),
        aaaamm="202606",
        operadora="Claro",
        tipo_contestacao=const.CENARIO_SEM_RETENCAO,
        tabelas_por_remuneracao=tabelas,
    )

    texto = "\n".join(p.text for p in documento.paragraphs)
    assert f"{const.CARTA_CIDADE_PADRAO}, 20 de Julho de 2026." in texto


def test_renderizar_carta_aceita_override_de_cidade():
    tabelas = gec.montar_tabelas_carta(
        pd.DataFrame([_linha_contest_completa("200", "011")])
    )

    documento = gec.renderizar_carta(
        numero_ct=1,
        data_carta=date(2026, 7, 20),
        aaaamm="202606",
        operadora="Claro",
        tipo_contestacao=const.CENARIO_SEM_RETENCAO,
        tabelas_por_remuneracao=tabelas,
        cidade="Rio de Janeiro",
    )

    texto = "\n".join(p.text for p in documento.paragraphs)
    assert "Rio de Janeiro, 20 de Julho de 2026." in texto


# --------------------------------------------------------------------------
# T-083 — gerar_arquivo_carta
# --------------------------------------------------------------------------
def test_gerar_arquivo_carta_grava_no_caminho_e_nome_corretos(tmp_path: Path):
    tabelas = gec.montar_tabelas_carta(
        pd.DataFrame([_linha_contest_completa("200", "011")])
    )
    documento = gec.renderizar_carta(
        numero_ct=363,
        data_carta=date(2026, 7, 20),
        aaaamm="202606",
        operadora="Claro",
        tipo_contestacao=const.CENARIO_SEM_RETENCAO,
        tabelas_por_remuneracao=tabelas,
        cidade="Rio de Janeiro",
    )

    caminho = gec.gerar_arquivo_carta(
        documento,
        operadora="Claro",
        numero_ct=363,
        aaaamm="202606",
        raiz_operadoras=tmp_path / "operadoras",
        raiz_controle_ct=tmp_path / "controle_ct",
    )

    assert caminho.name == "CT - 363.docx"
    assert caminho.parent == tmp_path / "operadoras" / "Claro" / "2026" / "202606" / "Contestações"
    assert caminho.is_file()


def test_gerar_arquivo_carta_salva_copia_em_controle_ct(tmp_path: Path):
    tabelas = gec.montar_tabelas_carta(
        pd.DataFrame([_linha_contest_completa("200", "011")])
    )
    documento = gec.renderizar_carta(
        numero_ct=363,
        data_carta=date(2026, 7, 20),
        aaaamm="202606",
        operadora="Claro",
        tipo_contestacao=const.CENARIO_SEM_RETENCAO,
        tabelas_por_remuneracao=tabelas,
        cidade="Rio de Janeiro",
    )

    gec.gerar_arquivo_carta(
        documento,
        operadora="Claro",
        numero_ct=363,
        aaaamm="202606",
        raiz_operadoras=tmp_path / "operadoras",
        raiz_controle_ct=tmp_path / "controle_ct",
    )

    caminho_copia = tmp_path / "controle_ct" / "2026" / "CT - 363.docx"
    assert caminho_copia.is_file()

    # A cópia deve ter o mesmo conteúdo textual do original.
    documento_copia = Document(str(caminho_copia))
    assert any("CT- 363/2026" in p.text for p in documento_copia.paragraphs)


# ---------------------------------------------------------------------------
# Q18 — a trava da numeração CT
# ---------------------------------------------------------------------------


class TestTravaDaNumeracao:
    """
    Q18, decidida em 2026-08-05.

    A seção crítica é o **par** ler-o-último → gravar-a-carta-com-ele. Dois
    processos entre esses dois passos leem o mesmo "último" e emitem o mesmo
    número — o que a decisão do cliente de 2026-07-31 proíbe.
    """

    def test_a_trava_existe_dentro_e_some_depois(self, tmp_path):
        with gec.travar_numeracao(tmp_path):
            assert (tmp_path / gec.ARQUIVO_TRAVA).is_file()

        assert not (tmp_path / gec.ARQUIVO_TRAVA).exists()

    def test_a_trava_e_liberada_mesmo_com_erro_dentro(self, tmp_path):
        with pytest.raises(RuntimeError):
            with gec.travar_numeracao(tmp_path):
                raise RuntimeError("falha ao renderizar a carta")

        assert not (tmp_path / gec.ARQUIVO_TRAVA).exists(), (
            "uma falha na emissão não pode deixar a pasta travada para o mês inteiro"
        )

    def test_trava_ja_presa_espera_e_segue_com_aviso(self, tmp_path, caplog):
        """
        Se o processo dono morrer sem liberar, o `.lock` fica para trás. Travar o
        mês inteiro por causa de um resíduo seria pior que o risco evitado —
        então espera o timeout, avisa alto e segue.
        """
        (tmp_path / gec.ARQUIVO_TRAVA).write_text("9999", encoding="utf-8")

        with caplog.at_level("WARNING"):
            with gec.travar_numeracao(tmp_path, timeout=0.1, intervalo=0.05):
                pass

        assert "trava" in caplog.text.lower()
        assert "duplicada" in caplog.text
        assert (tmp_path / gec.ARQUIVO_TRAVA).is_file(), (
            "quem não adquiriu a trava não pode removê-la do dono"
        )

    def test_a_trava_nao_entra_na_contagem_do_numero(self, tmp_path):
        """
        O `.lock` mora na mesma pasta que os CT. Se entrasse na varredura de
        nomes, a própria trava poderia virar número.
        """
        (tmp_path / "CT - 362.docx").write_text("x", encoding="utf-8")

        with gec.travar_numeracao(tmp_path):
            assert gec.obter_proximo_numero_carta(tmp_path) == 363

    def test_pasta_inexistente_e_numeracao_indeterminada(self, tmp_path):
        with pytest.raises(gec.NumeracaoCartaIndeterminada):
            with gec.travar_numeracao(tmp_path / "nao-existe"):
                pass
