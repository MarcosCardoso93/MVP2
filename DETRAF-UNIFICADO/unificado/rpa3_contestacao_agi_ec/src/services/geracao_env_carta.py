"""Serviço da HU-14 — arquivo `_ENV` + carta para a operadora.

**Escopo: T-080 (`_ENV`), T-081 (numeração CT), T-082 (renderizar carta), T-083 (salvar
carta).**

## Duas divergências deliberadas em relação à V2

A V2 descreve os dois artefatos saindo de modelos pré-existentes:

    ¶599 — "Ele copia o `Base_Contestação_<<nomedaoperadora>>_<<mesdodetraf>>_M`
            da operadora e salva na pasta com o novo nome: `..._ENV`."
    ¶601 — "O robô cria (…) a carta da operadora **a partir de um modelo
            pré-existente para cada operadora**."

Nenhum dos dois vale mais, por decisão do cliente:

- **`_ENV`** — não há modelo `_M`, porque a base de contestação deixou de ser
  arquivo e virou tabela (decisão de 2026-08-04). O `_ENV` é montado a partir do
  output da Consolidação (`montar_contest` + `consolidar_expectativa_vivo`);
- **Carta** — o modelo é **único para todas as operadoras** (decisão de
  2026-08-04), não um por operadora. Ela é montada do zero com `python-docx`, a
  partir de dois exemplos reais já emitidos.

⚠️ Isto é implementar a **decisão**, não o texto da V2 — e é deliberado. Um
relatório anterior chegou a afirmar que a V2 "não exigia modelo externo"; exigia.
O que mudou foi a regra, não a leitura do documento.

**T-082/T-083 (2026-07-27):** o modelo de carta foi substancialmente resolvido a partir
de dois exemplos reais já emitidos (`DOCS/Outros arquivos auxiliares/CT 251-2026...` e
`CT 252-2026...`, cenário SEM RETENÇÃO — ver D-3/D-13 em `TODO/decisoes.md`). Achado
importante: a tabela inserida na carta **é** quebrada por remuneração (corrige a nota
antiga de D-13, que assumia o contrário sem ver um exemplo real). **Assinatura (D-3,
resolvida em 2026-07-27):** os dois exemplos reais tinham signatários diferentes; o
usuário confirmou usar sempre a assinatura da carta real CT 252-2026 (Angélica) —
`ProvedorAssinaturaCartaPadrao`, sem mais placeholder.
"""

from __future__ import annotations

import os
import re
import shutil
import time
from contextlib import contextmanager
from datetime import date
from pathlib import Path
from typing import Callable, Optional

import pandas as pd
from docx import Document

from comum.config import configuration
from comum.config import constantes as const
from comum.config.logger_config import logger
from src.services import geracao_ext as gext
from src.services import mapa_remuneracao as mr
from comum.arquivos import estrutura_pastas as ep
from comum.arquivos import gerenciador as ga
from comum.arquivos import nomenclatura as nom
from comum.utils.decoradores import log_execucao

# Nomes das abas mantidas no _ENV (AI/09 §4.1: "manter apenas as abas Contest e TBRA").
ABA_CONTEST = const.ABA_CONTEST
ABA_TBRA = const.ABA_TBRA

#: Coluna interna com o sinal do analista. Existe só entre a seleção das linhas
#: contestadas e o consumo delas — **nunca chega ao `_ENV`**, que vai para a
#: operadora e não deve carregar campo de controle nosso.
COLUNA_SINAL = "_sinal"

# Colunas da tabela por remuneração inserida na carta (T-082) — sem
# `contestacao_a_enviar`/"Contestar" (AI/09 §4.2 item 4).
_COLUNAS_TABELA_CARTA = [
    "EOT",
    "minutos_tbra",
    "vb_tbra",
    "minutos_operadora",
    "vb_operadora",
    "minutos_diferenca",
    "vb_diferenca",
    "minutos_variacao_perc",
    "vb_variacao_perc",
]

_MES_EXTENSO = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
    7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro",
}

# Casa "CT 362", "CT-362", "CT - 362", "CT_362..." (com ou sem espaços/traço),
# case-insensitive — tolerante a variações de grafia observadas no nome real
# ("CT- 334/2025/WA*ID03", ver AI/09 §4.2 e o screenshot da carta CT-334/2025).
_PADRAO_NUMERO_CT = re.compile(r"CT[\s_-]*(\d+)", re.IGNORECASE)


def _extrair_numero_ct(nome_arquivo: str) -> int | None:
    """Extrai o número da carta de um nome de arquivo/pasta, ou `None` se não casar."""

    correspondencia = _PADRAO_NUMERO_CT.search(nome_arquivo)
    if not correspondencia:
        return None
    return int(correspondencia.group(1))


class NumeracaoCartaIndeterminada(RuntimeError):
    """Não foi possível determinar com segurança o próximo número de carta."""


#: Nome do arquivo de trava na pasta de controle CT.
ARQUIVO_TRAVA = ".numeracao-ct.lock"


@contextmanager
def travar_numeracao(
    pasta_controle_ct: Path, timeout: float = 60.0, intervalo: float = 0.5
):
    """
    Trava a pasta de controle CT enquanto um número é resolvido **e consumido**.

    **A seção crítica é o par, não cada metade** (Q18, decidida em 2026-08-05): o
    número sai do maior encontrado na pasta, e só passa a existir ali quando a
    carta é gravada. Dois processos entre esses dois passos leem o mesmo "último"
    e emitem cartas com o **mesmo número** — que é o que a decisão do cliente de
    2026-07-31 ("não pode haver duplicação") proíbe.

    Ficou mais necessária desde a Q25: a mesma execução passou a consumir **dois**
    números seguidos para uma operadora com cenário misto.

    Trava por arquivo, com `O_CREAT|O_EXCL` — atômico no Windows e no POSIX, e o
    único mecanismo que funciona quando a pasta é um compartilhamento de rede, que
    é o caso aqui (`\\\\lagoa\\...\\Correspondências Enviadas\\CT`).

    A trava é **liberada mesmo em erro**. Se o processo morrer sem liberar, o
    arquivo fica para trás e a próxima execução espera o timeout e segue com
    aviso — travar o mês inteiro por causa de um `.lock` órfão seria pior do que o
    risco que ele evita.

    Raises:
        NumeracaoCartaIndeterminada: Se a pasta não existir. Sem ela não há o que
            travar nem de onde ler o número.
    """

    pasta = Path(pasta_controle_ct)
    if not pasta.is_dir():
        raise NumeracaoCartaIndeterminada(
            f"[HU-14] Pasta de controle de numeração CT não encontrada: [{pasta}]."
        )

    trava = pasta / ARQUIVO_TRAVA
    descritor = None
    limite = time.monotonic() + timeout

    while True:
        try:
            descritor = os.open(trava, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
            break
        except FileExistsError:
            if time.monotonic() >= limite:
                logger.warning(
                    f"[HU-14] A trava [{trava.name}] continua presa após {timeout:.0f}s. "
                    f"Prosseguindo sem ela — provavelmente é resíduo de uma execução "
                    f"interrompida. Se houver outro robô emitindo carta agora, há "
                    f"risco de numeração duplicada."
                )
                break
            time.sleep(intervalo)

    try:
        if descritor is not None:
            os.write(descritor, f"{os.getpid()}".encode("utf-8"))
            os.close(descritor)
        yield
    finally:
        if descritor is not None:
            try:
                trava.unlink()
            except OSError as erro:
                logger.warning(
                    f"[HU-14] Não foi possível remover a trava [{trava}]: {erro}. "
                    f"A próxima execução vai esperar o timeout e seguir com aviso."
                )


@log_execucao
def obter_proximo_numero_carta(pasta_controle_ct: Path) -> int:
    """
    Determina o próximo número sequencial de carta CT (T-081).

    Lê os nomes de arquivo/pasta em `pasta_controle_ct` (tipicamente
    `estrutura_pastas.caminho_controle_ct(aaaamm)`), extrai o maior número `CT`
    encontrado e retorna o próximo (V2: "último `CT 362...` ⇒ novo `CT - 363...`").

    **Não pode haver duplicação de número.** Por decisão do cliente (2026-07-31),
    quando o último número não puder ser identificado — por qualquer motivo — o
    erro é acusado, em vez de assumir um valor. A versão anterior devolvia `1` em
    silêncio, o que reemitiria a carta nº 1 sobre uma sequência já existente
    sempre que a pasta estivesse inacessível.

    Exceção deliberada: em janeiro a pasta do ano é legitimamente nova e vazia.
    Para esse caso existe `CT_NUMERO_INICIAL` no `.env`, usado **somente** quando
    a pasta existe e não tem nenhuma carta. Sem a variável, erro — assim o
    bootstrap de uma sequência é sempre um ato deliberado.

    Args:
        pasta_controle_ct: Pasta `...\\Correspondências Enviadas\\CT\\{ano}`.

    Returns:
        O próximo número sequencial.

    Raises:
        NumeracaoCartaIndeterminada: Se a pasta não existir, não puder ser lida,
            ou estiver sem cartas e sem `CT_NUMERO_INICIAL` definido.
    """

    pasta_controle_ct = Path(pasta_controle_ct)

    if not pasta_controle_ct.is_dir():
        raise NumeracaoCartaIndeterminada(
            f"[HU-14] Pasta de controle de numeração CT não encontrada: "
            f"[{pasta_controle_ct}]. Sem ela não é possível saber o último número "
            f"emitido, e emitir uma carta agora arriscaria duplicar a numeração. "
            f"Verifique CAMINHO_CONTROLE_CT e o acesso à pasta."
        )

    try:
        nomes = [caminho.name for caminho in pasta_controle_ct.iterdir()]
    except OSError as erro:
        raise NumeracaoCartaIndeterminada(
            f"[HU-14] Não foi possível ler a pasta de controle de numeração CT "
            f"[{pasta_controle_ct}]: {erro}. Emitir uma carta agora arriscaria "
            f"duplicar a numeração."
        ) from erro

    numeros = [
        numero
        for nome in nomes
        if nome != ARQUIVO_TRAVA and (numero := _extrair_numero_ct(nome)) is not None
    ]

    if not numeros:
        if configuration.CT_NUMERO_INICIAL is not None:
            logger.warning(
                f"[HU-14] Nenhuma carta CT em [{pasta_controle_ct}] — iniciando a "
                f"sequência em {configuration.CT_NUMERO_INICIAL} "
                f"(CT_NUMERO_INICIAL)."
            )
            return configuration.CT_NUMERO_INICIAL

        raise NumeracaoCartaIndeterminada(
            f"[HU-14] A pasta [{pasta_controle_ct}] existe mas não contém nenhum "
            f"nome reconhecível como carta CT ({len(nomes)} item(ns) inspecionado"
            f"(s)). Não é possível determinar o último número emitido. Se esta é "
            f"de fato uma sequência nova (virada de ano), defina CT_NUMERO_INICIAL "
            f"no .env para iniciá-la deliberadamente."
        )

    ultimo = max(numeros)
    proximo = ultimo + 1
    logger.info(f"[HU-14] Último número CT encontrado: {ultimo}. Próximo: {proximo}.")
    return proximo


def nome_proxima_carta(pasta_controle_ct: Path) -> str:
    """Atalho: resolve o próximo número (T-081) e monta o nome `CT - {n}` (T-006)."""

    numero = obter_proximo_numero_carta(pasta_controle_ct)
    return nom.nome_carta(numero)


def selecionar_contestadas(
    df_contest: pd.DataFrame,
    referencia: str,
    obter_tipo_contestacao: Callable[[str, str, str, str, str], Optional[str]],
) -> pd.DataFrame:
    """
    Linhas efetivamente contestadas, com o sinal do analista em `COLUNA_SINAL`.

    **Gate de "linha contestada"** — o mesmo critério de `geracao_cont_proc`
    (D-13/T-100): ``contestacao_a_enviar == "S"`` **e** sinal do analista não
    `None` (cenário SEM ou COM retenção; nunca "sem contestação").

    Existe como função própria porque **dois consumidores precisam do mesmo
    conjunto**: `montar_abas_env`, que descarta o sinal antes de gravar o `_ENV`,
    e `separar_por_cenario`, que agrupa justamente por ele.
    """

    if df_contest.empty:
        return df_contest.copy()

    candidatas = df_contest[df_contest["contestacao_a_enviar"] == "S"].copy()
    if candidatas.empty:
        return candidatas

    candidatas[COLUNA_SINAL] = candidatas.apply(
        lambda linha: obter_tipo_contestacao(
            linha["eot_credora"], linha["eot_devedora"], referencia, linha["trafego"],
            linha["remuneracao"],
        ),
        axis=1,
    )
    return candidatas[candidatas[COLUNA_SINAL].notna()]


def separar_por_cenario(
    df_contest: pd.DataFrame,
    referencia: str,
    obter_tipo_contestacao: Callable[[str, str, str, str, str], Optional[str]],
) -> dict[str, pd.DataFrame]:
    """
    Agrupa as linhas contestadas por cenário — uma entrada por carta a emitir.

    **Q25, decidida em 2026-08-05: uma carta por cenário, cada uma com o seu
    número CT.** O sinal do analista é por chave
    ``(eot_operadora, eot_tbra, referência, tráfego, remuneração)``, então a mesma
    operadora pode ter linhas COM e SEM retenção no mesmo mês — e a carta é um
    documento com **um** texto de cenário.

    Antes desta decisão o código emitia uma carta só, com "prevalece COM
    retenção" — escolha nossa, registrada como pendência.

    Returns:
        ``{cenário: linhas daquele cenário}``, sem a coluna de sinal. Só entram
        cenários **presentes**: uma operadora só com retenção devolve uma entrada,
        não duas vazias. Dicionário vazio se nada foi contestado.
    """

    contestadas = selecionar_contestadas(
        df_contest, referencia, obter_tipo_contestacao
    )
    if contestadas.empty:
        return {}

    por_cenario: dict[str, pd.DataFrame] = {}
    for cenario in (const.CENARIO_COM_RETENCAO, const.CENARIO_SEM_RETENCAO):
        com_retencao = cenario == const.CENARIO_COM_RETENCAO
        linhas = contestadas[
            contestadas[COLUNA_SINAL].apply(gext.eh_com_retencao) == com_retencao
        ]
        if not linhas.empty:
            por_cenario[cenario] = linhas.drop(columns=[COLUNA_SINAL])

    return por_cenario


@log_execucao
def montar_abas_env(
    df_contest: pd.DataFrame,
    df_tbra_bruta: pd.DataFrame,
    referencia: str,
    indice_descritor: int,
    indice_remuneracao: dict[str, list[str]],
    obter_tipo_contestacao: Callable[[str, str, str, str, str], Optional[str]],
    indice_credora: int = const.COL_CREDORA,
    indice_devedora: int = const.COL_DEVEDORA,
    indice_trafego: int = const.COL_TRAFEGO,
) -> dict[str, pd.DataFrame]:
    """
    Monta as abas `Contest` e `TBRA` do `_ENV` (T-080), só com as linhas contestadas.

    AI/09 §4.1: "manter apenas as abas Contest e TBRA; apagar as demais" e "remover as
    linhas do que não será contestado — deixar somente os dados contestados". Como o
    modelo `_M` foi esclarecido como sendo o próprio output da Consolidação (D-3), as
    abas `RESUMO`/`{operadora}`/`DE PARA EOT` já são descartadas por não fazerem parte do
    conjunto de entrada desta função — não é necessário "apagá-las" explicitamente.

    **Gate de "linha contestada"** — mesmo critério de `geracao_cont_proc` (D-13/T-100):
    `contestacao_a_enviar == "S"` **e** sinal do analista não `None` (cenário SEM/COM
    retenção; nunca "sem contestação").

    Args:
        df_contest: Saída de `consolidacao_contestacao.montar_contest` (T-023).
        df_tbra_bruta: Aba `TBRA` consolidada (T-021), linha a linha (sem agregação),
            até `R$_Bruto`.
        referencia: Mês de referência do Detraf da contestação (`AAAAMM`).
        indice_descritor: Índice (0-based) da coluna de descritor em `df_tbra_bruta` —
            sem posição fixa na documentação, informado pelo chamador.
        indice_remuneracao: Índice `descritor final -> [remunerações]` (D-5).
        obter_tipo_contestacao: Callable injetada (repository via controller) que resolve
            o sinal COM/SEM retenção por
            `(eot_operadora, eot_tbra, referencia, trafego, remuneracao)` (T-024/D-6; a
            remuneração entra na chave desde 2026-07-28).
        indice_credora, indice_devedora, indice_trafego: Índices documentados (D-8) usados
            para casar cada linha bruta da `TBRA` com a chave contestada do `Contest`.

    Returns:
        Dicionário ``{"Contest": DataFrame, "TBRA": DataFrame}``. Ambos vazios se nada
        for contestado.
    """

    if df_contest.empty:
        return {ABA_CONTEST: df_contest.copy(), ABA_TBRA: df_tbra_bruta.iloc[0:0].copy()}

    contestadas = selecionar_contestadas(
        df_contest, referencia, obter_tipo_contestacao
    ).drop(columns=[COLUNA_SINAL], errors="ignore")

    if contestadas.empty:
        logger.info("[HU-14 _ENV] Nenhuma linha contestada — abas ficam vazias.")
        return {
            ABA_CONTEST: df_contest.iloc[0:0].copy(),
            ABA_TBRA: df_tbra_bruta.iloc[0:0].copy(),
        }

    chaves_contestadas = set(
        zip(
            contestadas["eot_credora"],
            contestadas["eot_devedora"],
            contestadas["remuneracao"],
            contestadas["trafego"],
        )
    )

    tbra = df_tbra_bruta.copy()
    tbra["_remuneracao"] = mr.resolver_series(
        tbra.iloc[:, indice_descritor], indice_remuneracao
    )
    tbra["_chave"] = list(
        zip(
            tbra.iloc[:, indice_credora].astype(str).str.strip(),
            tbra.iloc[:, indice_devedora].astype(str).str.strip(),
            tbra["_remuneracao"],
            tbra.iloc[:, indice_trafego].astype(str).str.strip(),
        )
    )
    tbra_filtrada = tbra[tbra["_chave"].isin(chaves_contestadas)].drop(
        columns=["_remuneracao", "_chave"]
    )

    logger.info(
        f"[HU-14 _ENV] Contest: {contestadas.shape[0]} linha(s) contestada(s); "
        f"TBRA: {tbra_filtrada.shape[0]} linha(s) correspondente(s)."
    )
    return {ABA_CONTEST: contestadas, ABA_TBRA: tbra_filtrada}


@log_execucao
def gerar_arquivo_env(
    abas_env: dict[str, pd.DataFrame],
    operadora: str,
    aaaamm: str,
    raiz_operadoras: Path | None = None,
) -> Optional[Path]:
    """
    Grava o arquivo `_ENV` na pasta `Contestações` (T-080) — **só se houver contestação**.

    Se ambas as abas estiverem vazias (nenhuma linha contestada), **nenhum arquivo é
    criado** — mesmo critério de `gerar_arquivo_int`/`gerar_arquivo_cont_proc`: o `_ENV`
    só existe nos cenários SEM/COM retenção (AI/09 §0, matriz).

    Args:
        abas_env: Saída de :func:`montar_abas_env` (``{"Contest": df, "TBRA": df}``).
        operadora: Nome da operadora (usado no nome do arquivo e na pasta).
        aaaamm: Período de referência (`AAAAMM`).
        raiz_operadoras: Raiz `...\\Operadoras`. Default: `CAMINHO_OPERADORAS` (config).

    Returns:
        Caminho do arquivo `.xlsx` gravado, ou ``None`` se nada foi contestado.
    """

    if all(df.empty for df in abas_env.values()):
        logger.info("[HU-14 _ENV] Nada a gravar (nenhuma linha contestada).")
        return None

    pasta_contestacoes = ep.caminho_contestacoes(
        operadora, aaaamm, raiz_operadoras=raiz_operadoras, criar=True
    )
    nome_arquivo = nom.nome_env(operadora, aaaamm)
    caminho = pasta_contestacoes / f"{nome_arquivo}{const.EXTENSAO_EXCEL}"

    ga.salvar_planilhas(abas_env, caminho, incluir_cabecalho=True)

    logger.info(f"[HU-14 _ENV] Arquivo gravado em [{caminho}].")
    return caminho


class ProvedorAssinaturaCarta:
    """Contrato para resolver o signatário (nome, cargo) da carta.

    Isolado atrás desta interface porque os dois exemplos reais (CT 251/CT 252-2026,
    mesmo dia) tinham signatários diferentes com o mesmo cargo genérico — a escolha não
    era óbvia até o usuário decidir (D-3, resolvida em 2026-07-27).
    """

    def resolver(self, operadora: str) -> tuple[str, str]:
        raise NotImplementedError


class ProvedorAssinaturaCartaPadrao(ProvedorAssinaturaCarta):
    """Assinatura fixa confirmada pelo usuário (D-3, 2026-07-27).

    Usa sempre a assinatura da carta real
    ``CT 252-2026-DIR-A1-tbrasilxserctel-Contestação_202606.docx`` (Angélica), para
    qualquer operadora — não há distinção por operadora/região.
    """

    def resolver(self, operadora: str) -> tuple[str, str]:
        return (const.CARTA_ASSINATURA_NOME, const.CARTA_ASSINATURA_CARGO)


@log_execucao
def montar_tabelas_carta(df_contest_contestado: pd.DataFrame) -> dict[str, dict]:
    """
    Monta as tabelas por remuneração inseridas no corpo da carta (T-082).

    **Granularidade (corrige D-13, 2026-07-27):** os exemplos reais (CT 251/252-2026)
    confirmam que a tabela da carta é quebrada **por remuneração** (blocos "TU-RL",
    "VU-M" etc., cada um com sub-rótulo de tipo de serviço quando disponível) — não há
    colapso por EOT×Tráfego como se assumia antes. Usa a aba `Contest` já filtrada por
    `montar_abas_env` (T-080) diretamente, sem agregação adicional.

    Args:
        df_contest_contestado: Aba `Contest` já filtrada (saída de
            ``montar_abas_env(...)["Contest"]``, T-080) — colunas `eot_credora`,
            `eot_devedora`, `remuneracao`, `minutos_tbra`, `vb_tbra`,
            `minutos_operadora`, `vb_operadora`, `minutos_diferenca`, `vb_diferenca`,
            `minutos_variacao_perc`, `vb_variacao_perc` e, opcionalmente,
            `tipo_operacao` (rótulo SMS/STFC visto nos exemplos reais).

    Returns:
        Dicionário ``{remuneracao: {"tipo_operacao": str | None, "tabela": DataFrame}}``,
        na ordem em que a remuneração aparece em ``df_contest_contestado``. Cada
        `DataFrame` traz a coluna `EOT` (formato `devedora/credora`, D-11) + as colunas
        numéricas, com uma linha `TOTAL` ao final (variação percentual em branco na
        linha `TOTAL` — mesmo padrão visto nos exemplos reais). **Sem** a coluna
        `contestacao_a_enviar`/"Contestar" (AI/09 §4.2). Vazio se não houver nada
        contestado.
    """

    if df_contest_contestado.empty:
        return {}

    df = df_contest_contestado.copy()
    df["EOT"] = (
        df["eot_devedora"].astype(str) + "/" + df["eot_credora"].astype(str)
    )

    resultado: dict[str, dict] = {}
    for remuneracao, grupo in df.groupby("remuneracao", sort=False):
        tipo_operacao = None
        if "tipo_operacao" in grupo.columns:
            valores = grupo["tipo_operacao"].dropna().unique()
            if len(valores) == 1:
                tipo_operacao = valores[0]

        tabela = grupo[_COLUNAS_TABELA_CARTA].reset_index(drop=True)

        totais = tabela.drop(columns="EOT").sum(numeric_only=True).to_dict()
        linha_total = {"EOT": "TOTAL", **totais}
        # Variação percentual não é somável — linha TOTAL fica em branco (mesmo padrão
        # visto nos exemplos reais CT 251/252-2026).
        linha_total["minutos_variacao_perc"] = None
        linha_total["vb_variacao_perc"] = None

        tabela = pd.concat([tabela, pd.DataFrame([linha_total])], ignore_index=True)
        resultado[str(remuneracao)] = {"tipo_operacao": tipo_operacao, "tabela": tabela}

    logger.info(f"[HU-14 Carta] {len(resultado)} tabela(s) de remuneração montada(s).")
    return resultado


def _inserir_tabela(documento: Document, tabela: pd.DataFrame) -> None:
    """Insere `tabela` como uma tabela nativa do Word em `documento` (T-082)."""

    colunas = list(tabela.columns)
    tabela_docx = documento.add_table(rows=1, cols=len(colunas))
    for indice, coluna in enumerate(colunas):
        tabela_docx.rows[0].cells[indice].text = coluna

    for _, linha in tabela.iterrows():
        celulas = tabela_docx.add_row().cells
        for indice, coluna in enumerate(colunas):
            valor = linha[coluna]
            celulas[indice].text = "" if pd.isna(valor) else str(valor)


@log_execucao
def renderizar_carta(
    numero_ct: int,
    data_carta: date,
    aaaamm: str,
    operadora: str,
    tipo_contestacao: str,
    tabelas_por_remuneracao: dict[str, dict],
    cidade: Optional[str] = None,
    provedor_assinatura: Optional[ProvedorAssinaturaCarta] = None,
) -> Document:
    """
    Renderiza o `.docx` da carta de contestação (T-082), a partir dos exemplos reais
    CT 251/252-2026 (D-3).

    **Não replica** logo/rodapé/fontes do arquivo real — mesmo gap aceito em D-12
    (EXT/INT): conteúdo é o critério de aceite testável, formatação exata fica para
    quando houver necessidade real.

    Args:
        numero_ct: Número sequencial da carta (T-081, `obter_proximo_numero_carta`).
        data_carta: Data de emissão (injetada pelo chamador — nunca `datetime.now()`
            dentro do service, para manter a função determinística/testável).
        aaaamm: Mês de referência do Detraf (usado no "Assunto:" e no corpo).
        operadora: Nome da operadora (usado no corpo).
        tipo_contestacao: `constantes_epico4.CENARIO_SEM_RETENCAO` ou
            `CENARIO_COM_RETENCAO` (texto COM retenção é assumido por substituição
            simples da frase SEM retenção — sem exemplo real ainda, ver D-3).
        tabelas_por_remuneracao: Saída de :func:`montar_tabelas_carta`.
        cidade: Cidade da data de emissão. **RESOLVIDO (D-3, 2026-07-27):** usuário
            confirmou `constantes_epico4.CARTA_CIDADE_PADRAO` ("São Paulo") como padrão
            fixo — os 2 exemplos reais divergiam no mesmo dia (Rio de Janeiro/São Paulo).
            Parâmetro mantido para permitir override pontual, se necessário.
        provedor_assinatura: Resolve `(nome, cargo)` do signatário; default
            :class:`ProvedorAssinaturaCartaPadrao` (assinatura fixa confirmada em D-3).

    Returns:
        Documento `python-docx` pronto para ser salvo (:func:`gerar_arquivo_carta`).
    """

    if cidade is None:
        cidade = const.CARTA_CIDADE_PADRAO

    documento = Document()

    documento.add_paragraph(f"{const.PREFIXO_CARTA}- {numero_ct}/{data_carta.year}")
    documento.add_paragraph(
        f"{cidade}, {data_carta.day} de {_MES_EXTENSO[data_carta.month]} de "
        f"{data_carta.year}."
    )
    documento.add_paragraph(
        f"ASSUNTO: {const.CARTA_ASSUNTO_TEMPLATE.format(aaaamm=aaaamm)}"
    )
    documento.add_paragraph(const.CARTA_SAUDACAO)
    documento.add_paragraph(
        const.CARTA_CORPO_TEMPLATE.format(
            tipo_contestacao=tipo_contestacao,
            aaaamm=aaaamm,
            operadora=operadora.upper(),
        )
    )

    for remuneracao, dados in tabelas_por_remuneracao.items():
        titulo = remuneracao
        if dados.get("tipo_operacao"):
            titulo = f"{remuneracao}    {dados['tipo_operacao']}"
        documento.add_paragraph(titulo)
        _inserir_tabela(documento, dados["tabela"])

    documento.add_paragraph(const.CARTA_FECHO)

    provedor = provedor_assinatura or ProvedorAssinaturaCartaPadrao()
    nome, cargo = provedor.resolver(operadora)
    documento.add_paragraph(nome)
    documento.add_paragraph(cargo)

    logger.info(
        f"[HU-14 Carta] Carta CT-{numero_ct} renderizada "
        f"({len(tabelas_por_remuneracao)} tabela(s))."
    )
    return documento


@log_execucao
def gerar_arquivo_carta(
    documento: Document,
    operadora: str,
    numero_ct: int,
    aaaamm: str,
    raiz_operadoras: Path | None = None,
    raiz_controle_ct: Path | None = None,
) -> Path:
    """
    Salva a carta (T-083) em `Contestações` e uma cópia em `...\\CT\\{ano}` (AI/09 §4.2
    passo 5).

    Args:
        documento: Saída de :func:`renderizar_carta`.
        operadora: Nome da operadora (usado na pasta).
        numero_ct: Número sequencial da carta (usado no nome do arquivo).
        aaaamm: Período de referência (usado na pasta/ano de controle).
        raiz_operadoras: Raiz `...\\Operadoras`. Default: `CAMINHO_OPERADORAS` (config).
        raiz_controle_ct: Raiz `...\\Correspondências Enviadas\\CT`. Default:
            `CAMINHO_CONTROLE_CT` (config).

    Returns:
        Caminho do arquivo `.docx` gravado em `Contestações`.
    """

    pasta_contestacoes = ep.caminho_contestacoes(
        operadora, aaaamm, raiz_operadoras=raiz_operadoras, criar=True
    )
    nome_arquivo = f"{nom.nome_carta(numero_ct)}{const.EXTENSAO_DOCX}"
    caminho = pasta_contestacoes / nome_arquivo
    documento.save(str(caminho))

    pasta_controle = ep.caminho_controle_ct(
        aaaamm, raiz_controle_ct=raiz_controle_ct, criar=True
    )
    caminho_copia = pasta_controle / nome_arquivo
    shutil.copy2(caminho, caminho_copia)

    logger.info(
        f"[HU-14 Carta] Carta gravada em [{caminho}], cópia em [{caminho_copia}]."
    )
    return caminho
